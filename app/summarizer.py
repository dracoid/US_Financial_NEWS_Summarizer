import os
import time
import pandas as pd
import warnings
from newspaper import Article, Config
from transformers import pipeline
from app.db import insert_summary, summary_exists
from app.telegram_sender import send_docx_files
from app.config import EXCEL_PATH, DOCX_SAVE_DIR
from docx import Document
from datetime import datetime

# NumPy 경고 무시 (PyTorch 내부에서 발생하는 NumPy 관련 경고 억제)
warnings.filterwarnings("ignore", message="Failed to initialize NumPy")

# 사용자 User-Agent 설정 (크롤링 차단 회피용)
user_config = Config()
user_config.browser_user_agent = (
    "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 "
    "(KHTML, like Gecko) Chrome/123.0.0.0 Safari/537.36"
)

# Hugging Face summarizer 로딩
print("Device set to use cpu")
summarizer = pipeline("summarization", model="facebook/bart-large-cnn", device=0)

def summarize_and_store():
    print("🛠 DB 초기화...")

    # Excel 파일 로드
    df = pd.read_excel(EXCEL_PATH)
    df = df.sort_values(by="published_dt")

    print("📰 뉴스 요약 시작...")

    # 티커별 그룹핑
    grouped = df.groupby("ticker")

    for ticker, group in grouped:
        doc = Document()
        doc.add_heading(f"News Summary for {ticker}", 0)
        saved = False

        for _, row in group.iterrows():
            title = row["title"]
            link = row["link"]
            date = row["published_dt"]

            # 중복 여부 확인
            if summary_exists(title):
                print(f"[SKIP] 이미 저장됨: {title}")
                continue

            try:
                article = Article(link, config=user_config)
                article.download()
                article.parse()
                text = article.text

                # 너무 짧은 기사 제외
                if len(text) < 200:
                    print(f"[SKIP] 기사 너무 짧음: {title}")
                    continue

                # 요약 수행
                summary = summarizer(text[:1024], max_length=130, min_length=30, do_sample=False)[0]['summary_text']

                # DB 저장
                insert_summary(title, date, summary, link)

                # DOCX 저장 준비
                doc.add_heading(str(date), level=2)
                doc.add_paragraph(f"🔗 {link}")
                doc.add_paragraph(f"📰 {title}")
                doc.add_paragraph(summary)
                doc.add_paragraph("-" * 30)
                print(f"[✓] {ticker}: {title}")
                saved = True
                time.sleep(1)

            except Exception as e:
                insert_summary(title, date, f"(크롤링 실패: {e})", link)
                print(f"[Error] {link}: {e}")
                continue

        # DOCX 저장
        if saved:
            filename = os.path.join(DOCX_SAVE_DIR, f"{ticker}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx")
            doc.save(filename)
            print(f"[📄 저장 완료] {filename}")

    # 텔레그램 전송
    send_docx_files()

