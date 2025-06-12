# app/exporter.py

from docx import Document
from .db import get_connection
from .config import DOCX_SAVE_DIR
import os

def export_docx_all():
    conn = get_connection()
    cursor = conn.cursor()

    cursor.execute("SELECT DISTINCT ticker FROM news_summaries")
    tickers = [r[0] for r in cursor.fetchall()]

    for ticker in tickers:
        cursor.execute("""
        SELECT published_dt, title, link, summary 
        FROM news_summaries 
        WHERE ticker = ? ORDER BY published_dt ASC
        """, (ticker,))
        rows = cursor.fetchall()

        if not rows:
            continue

        doc = Document()
        doc.add_heading(f"News Summary for {ticker}", 0)

        for date, title, link, summary in rows:
            doc.add_heading(str(date), level=2)
            doc.add_paragraph(f"🔗 {link}")
            doc.add_paragraph(f"📰 {title}")
            doc.add_paragraph(summary)
            doc.add_paragraph("-" * 30)

        save_path = os.path.join(DOCX_SAVE_DIR, f"{ticker.replace('^','')}_summary.docx")
        doc.save(save_path)
        print(f"📄 저장 완료: {save_path}")

    conn.close()
